"""
Word报告生成模块
用于将股票分析结果生成Word文档
"""
import os
import pandas as pd
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO
import matplotlib.pyplot as plt
import numpy as np
import akshare as ak


def create_word_report(stock_data, analysis_report, charts_data, user_question, symbol, start_date=None, end_date=None):
    """创建Word报告文档"""
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(f'{stock_data["实时行情"]["名称"]} ({symbol}) 股票分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分析时间
        doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'分析区间: {start_date if start_date else "最近"} 至 {end_date if end_date else "当前"}')
        doc.add_paragraph(f'用户问题: {user_question}')
        
        # 添加分页
        doc.add_page_break()
        
        # 添加股票基本信息
        doc.add_heading('股票基本信息', level=1)
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        # 添加基本信息行
        info_data = [
            ['股票名称', stock_data["实时行情"]["名称"]],
            ['股票代码', symbol],
            ['当前价格', f'{stock_data["实时行情"]["最新价"]}元'],
            ['涨跌幅', f'{stock_data["实时行情"]["涨跌幅"]}%'],
            ['成交量', f'{stock_data["实时行情"]["成交量"]:,}'],
            ['成交额', f'{stock_data["实时行情"]["成交额"]:,}'],
            ['振幅', f'{stock_data["实时行情"]["振幅"]}%'],
            ['换手率', f'{stock_data["实时行情"]["换手率"]}%'],
        ]
        
        for info in info_data:
            row_cells = info_table.add_row().cells
            row_cells[0].text = info[0]
            row_cells[1].text = str(info[1])
        
        # 添加技术指标
        doc.add_heading('技术指标分析', level=1)
        # 这里需要从stock_data中提取技术指标，如果存在的话
        # 为了简化，我们暂时跳过这部分，实际应用中需要从stock_data中获取
        
        # 添加AI分析报告
        doc.add_heading('AI分析报告', level=1)
        doc.add_paragraph(analysis_report)
        
        # 添加图表
        doc.add_heading('股价走势图', level=1)
        if charts_data:
            doc.add_paragraph('以下是该股票的价格走势和成交量图:')
            
            # 添加趋势图
            if 'trend' in charts_data and charts_data['trend']:
                doc.add_heading('价格趋势图', level=2)
                doc.add_paragraph('显示股票价格的走势及主要移动平均线:')
                # 保存图表到临时文件并添加到文档
                chart_path = charts_data['trend']
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            # 添加成交量图
            if 'volume' in charts_data and charts_data['volume']:
                doc.add_heading('成交量图', level=2)
                doc.add_paragraph('显示股票成交量的变化:')
                chart_path = charts_data['volume']
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            # 添加技术指标图
            if 'rsi' in charts_data and charts_data['rsi']:
                doc.add_heading('技术指标图', level=2)
                doc.add_paragraph('显示RSI等技术指标:')
                chart_path = charts_data['rsi']
                if os.path.exists(chart_path):
                    doc.add_picture(chart_path, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加历史数据表格
        doc.add_heading('历史数据', level=1)
        if stock_data.get("历史数据"):
            hist_data = stock_data["历史数据"][-10:]  # 只取最近10天数据
            if hist_data:
                hist_table = doc.add_table(rows=1, cols=6)
                hist_table.style = 'Table Grid'
                hdr_cells = hist_table.rows[0].cells
                hdr_cells[0].text = '日期'
                hdr_cells[1].text = '开盘'
                hdr_cells[2].text = '收盘'
                hdr_cells[3].text = '最高'
                hdr_cells[4].text = '最低'
                hdr_cells[5].text = '成交量'
                
                for record in hist_data:
                    row_cells = hist_table.add_row().cells
                    row_cells[0].text = str(record.get('日期', ''))
                    row_cells[1].text = f"{record.get('开盘', 0):.2f}"
                    row_cells[2].text = f"{record.get('收盘', 0):.2f}"
                    row_cells[3].text = f"{record.get('最高', 0):.2f}"
                    row_cells[4].text = f"{record.get('最低', 0):.2f}"
                    row_cells[5].text = f"{record.get('成交量', 0):,}"
        
        # 添加风险提示
        doc.add_heading('风险提示', level=1)
        risk_paragraph = doc.add_paragraph()
        risk_paragraph.add_run('投资有风险，决策需谨慎。以上分析仅供参考，不构成投资建议。').bold = True
        risk_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{symbol}_analysis_report_{timestamp}.docx"
        
        # 保存文档
        doc.save(filename)
        
        return filename
    except Exception as e:
        print(f"⚠️ 生成Word报告时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def create_word_report_in_memory(stock_data, analysis_report, charts_data, user_question, symbol, start_date=None, end_date=None):
    """在内存中创建Word报告文档，不保存到文件系统"""
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(f'{stock_data["实时行情"]["名称"]} ({symbol}) 股票分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加分析时间
        doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'分析区间: {start_date if start_date else "最近"} 至 {end_date if end_date else "当前"}')
        doc.add_paragraph(f'用户问题: {user_question}')
        
        # 添加分页
        doc.add_page_break()
        
        # 添加股票基本信息
        doc.add_heading('股票基本信息', level=1)
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'
        
        # 添加基本信息行
        info_data = [
            ['股票名称', stock_data["实时行情"]["名称"]],
            ['股票代码', symbol],
            ['当前价格', f'{stock_data["实时行情"]["最新价"]}元'],
            ['涨跌幅', f'{stock_data["实时行情"]["涨跌幅"]}%'],
            ['成交量', f'{stock_data["实时行情"]["成交量"]:,}'],
            ['成交额', f'{stock_data["实时行情"]["成交额"]:,}'],
            ['振幅', f'{stock_data["实时行情"]["振幅"]}%'],
            ['换手率', f'{stock_data["实时行情"]["换手率"]}%'],
        ]
        
        for info in info_data:
            row_cells = info_table.add_row().cells
            row_cells[0].text = info[0]
            row_cells[1].text = str(info[1])
        
        # 添加技术指标
        doc.add_heading('技术指标分析', level=1)
        # 这里需要从stock_data中提取技术指标，如果存在的话
        # 为了简化，我们暂时跳过这部分，实际应用中需要从stock_data中获取
        
        # 添加AI分析报告
        doc.add_heading('AI分析报告', level=1)
        doc.add_paragraph(analysis_report)
        
        # 添加图表（如果charts_data包含内存中的图片）
        doc.add_heading('股价走势图', level=1)
        if charts_data:
            doc.add_paragraph('以下是该股票的价格走势和成交量图:')
            
            # 添加趋势图
            if 'trend' in charts_data and charts_data['trend']:
                doc.add_heading('价格趋势图', level=2)
                doc.add_paragraph('显示股票价格的走势及主要移动平均线:')
                # 如果是文件路径
                if isinstance(charts_data['trend'], str) and os.path.exists(charts_data['trend']):
                    try:
                        doc.add_picture(charts_data['trend'], width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"添加趋势图失败: {e}")
                        doc.add_paragraph('趋势图: 无法加载图片')
                # 如果是BytesIO对象
                elif hasattr(charts_data['trend'], 'getvalue'):
                    try:
                        # 将BytesIO保存为临时文件再添加到文档
                        temp_path = f"temp_trend_chart_{int(time.time())}.png"
                        with open(temp_path, 'wb') as f:
                            f.write(charts_data['trend'].getvalue())
                        doc.add_picture(temp_path, width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 删除临时文件
                        os.remove(temp_path)
                    except Exception as e:
                        print(f"添加趋势图失败: {e}")
                        doc.add_paragraph('趋势图: 无法加载图片')
                else:
                    doc.add_paragraph('趋势图: 无数据')
        
            # 添加成交量图
            if 'volume' in charts_data and charts_data['volume']:
                doc.add_heading('成交量图', level=2)
                doc.add_paragraph('显示股票成交量的变化:')
                # 如果是文件路径
                if isinstance(charts_data['volume'], str) and os.path.exists(charts_data['volume']):
                    try:
                        doc.add_picture(charts_data['volume'], width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"添加成交量图失败: {e}")
                        doc.add_paragraph('成交量图: 无法加载图片')
                # 如果是BytesIO对象
                elif hasattr(charts_data['volume'], 'getvalue'):
                    try:
                        # 将BytesIO保存为临时文件再添加到文档
                        temp_path = f"temp_volume_chart_{int(time.time())}.png"
                        with open(temp_path, 'wb') as f:
                            f.write(charts_data['volume'].getvalue())
                        doc.add_picture(temp_path, width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 删除临时文件
                        os.remove(temp_path)
                    except Exception as e:
                        print(f"添加成交量图失败: {e}")
                        doc.add_paragraph('成交量图: 无法加载图片')
                else:
                    doc.add_paragraph('成交量图: 无数据')
        
            # 添加技术指标图
            if 'rsi' in charts_data and charts_data['rsi']:
                doc.add_heading('技术指标图', level=2)
                doc.add_paragraph('显示RSI等技术指标:')
                # 如果是文件路径
                if isinstance(charts_data['rsi'], str) and os.path.exists(charts_data['rsi']):
                    try:
                        doc.add_picture(charts_data['rsi'], width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"添加RSI图失败: {e}")
                        doc.add_paragraph('RSI图: 无法加载图片')
                # 如果是BytesIO对象
                elif hasattr(charts_data['rsi'], 'getvalue'):
                    try:
                        # 将BytesIO保存为临时文件再添加到文档
                        temp_path = f"temp_rsi_chart_{int(time.time())}.png"
                        with open(temp_path, 'wb') as f:
                            f.write(charts_data['rsi'].getvalue())
                        doc.add_picture(temp_path, width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 删除临时文件
                        os.remove(temp_path)
                    except Exception as e:
                        print(f"添加RSI图失败: {e}")
                        doc.add_paragraph('RSI图: 无法加载图片')
                else:
                    doc.add_paragraph('RSI图: 无数据')
        
        # 添加历史数据表格
        doc.add_heading('历史数据', level=1)
        if stock_data.get("历史数据"):
            hist_data = stock_data["历史数据"][-10:]  # 只取最近10天数据
            if hist_data:
                hist_table = doc.add_table(rows=1, cols=6)
                hist_table.style = 'Table Grid'
                hdr_cells = hist_table.rows[0].cells
                hdr_cells[0].text = '日期'
                hdr_cells[1].text = '开盘'
                hdr_cells[2].text = '收盘'
                hdr_cells[3].text = '最高'
                hdr_cells[4].text = '最低'
                hdr_cells[5].text = '成交量'
                
                for record in hist_data:
                    row_cells = hist_table.add_row().cells
                    row_cells[0].text = str(record.get('日期', ''))
                    row_cells[1].text = f"{record.get('开盘', 0):.2f}"
                    row_cells[2].text = f"{record.get('收盘', 0):.2f}"
                    row_cells[3].text = f"{record.get('最高', 0):.2f}"
                    row_cells[4].text = f"{record.get('最低', 0):.2f}"
                    row_cells[5].text = f"{record.get('成交量', 0):,}"
        
        # 添加风险提示
        doc.add_heading('风险提示', level=1)
        risk_paragraph = doc.add_paragraph()
        risk_paragraph.add_run('投资有风险，决策需谨慎。以上分析仅供参考，不构成投资建议。').bold = True
        risk_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 将文档保存到内存中的BytesIO对象
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        print(f"⚠️ 生成Word报告时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return None