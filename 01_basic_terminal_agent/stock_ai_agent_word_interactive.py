#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
股票分析AI智能体 - 交互式Word报告生成版本
功能：交互式持续对话，可切换股票，持续提问，并生成包含图表的Word文档

依赖库安装命令：
    pip install akshare pandas numpy matplotlib python-docx
    或使用特定Python版本：
    py -你的python版本 -m pip install akshare pandas numpy matplotlib python-docx

使用方法：
    (可以直接运行)python stock_ai_agent_word_interactive.py --interactive  # 交互模式（可连续提问并生成Word报告）

API配置说明：
    请将 QWEN_API_KEY 替换为您的真实通义千问API密钥

文件生成逻辑：
    1. Word报告：每次提问后都会生成一个以股票代码和时间戳命名的docx文件
    2. 报告文件命名格式：{股票代码}_analysis_report_{YYYYMMDD_HHMM}.docx
    3. 文件保存路径：与脚本相同的目录
"""

import argparse
import json
import akshare as ak
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import matplotlib
matplotlib.use('Agg')  # 使用非GUI后端
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

# ==================== 配置区 ====================
# 在这里配置你的大模型API（以通义千问API为例）
QWEN_API_KEY = "sk-ed31504ad0554161a6191b41287b9c88"  # 替换为你的真实API密钥
QWEN_API_URL = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"  # 通义千问API地址

def get_stock_data(symbol, period="近期"):
    """
    使用AKShare获取股票数据
    支持输入：'000001'（平安银行）、'000001.SZ'、'sh600000'等格式
    """
    try:
        # 获取实时行情
        spot_df = ak.stock_zh_a_spot_em()
        stock_info = spot_df[spot_df['代码'] == symbol.split('.')[0]]

        if stock_info.empty:
            return None, "未找到该股票代码"

        # 获取历史K线数据（近期）
        end_date = datetime.now().strftime("%Y%m%d")
        start_date = (datetime.now() - timedelta(days=30)).strftime("%Y%m%d")
        hist_df = ak.stock_zh_a_hist(symbol=symbol, period="daily", start_date=start_date, end_date=end_date, adjust="qfq")

        # 跳过基本面数据获取，因为该接口经常超时
        print("⚠️ 提示: 为提高稳定性，已跳过基本面数据获取")
        fundamentals = "无数据"

        # 整理数据
        # 将历史数据中的日期等非JSON序列化类型转换为字符串
        hist_data_serializable = []
        for record in hist_df.tail(5).to_dict('records'):
            serializable_record = {}
            for k, v in record.items():
                # 检查是否为时间戳或日期类型
                if hasattr(v, 'strftime') or str(type(v)) in ["<class 'pandas._libs.tslibs.timestamps.Timestamp'>", "<class 'datetime.date'>", "<class 'numpy.datetime64'>"]:
                    serializable_record[k] = str(v)
                elif pd.isna(v):
                    serializable_record[k] = None
                else:
                    serializable_record[k] = v
            hist_data_serializable.append(serializable_record)
            
        data = {
            "实时行情": {
                "名称": stock_info.iloc[0]['名称'],
                "最新价": float(stock_info.iloc[0]['最新价']),
                "涨跌幅": float(str(stock_info.iloc[0]['涨跌幅']).rstrip('%')) if isinstance(stock_info.iloc[0]['涨跌幅'], str) else float(stock_info.iloc[0]['涨跌幅']),
                "成交量": int(float(stock_info.iloc[0]['成交量'])) if not pd.isna(stock_info.iloc[0]['成交量']) else 0,
                "成交额": int(float(stock_info.iloc[0]['成交额'])) if not pd.isna(stock_info.iloc[0]['成交额']) else 0,
            },
            "历史数据": hist_data_serializable,  # 最近5天，已处理为可序列化格式
            "基本面": fundamentals
        }
        return data, "成功"
    except Exception as e:
        return None, f"获取数据时出错: {str(e)}"

def calculate_technical_indicators(hist_data):
    """计算基础技术指标"""
    if not hist_data or len(hist_data) < 10:
        return {}

    # 转换日期字段为字符串，以便JSON序列化
    df = pd.DataFrame([{k: str(v) if isinstance(v, (pd.Timestamp, datetime.date)) else v for k, v in record.items()} for record in hist_data])
    close_prices = df['收盘'].astype(float).values

    # 简单移动平均
    sma_5 = np.mean(close_prices[-5:]) if len(close_prices) >= 5 else None
    sma_10 = np.mean(close_prices[-10:]) if len(close_prices) >= 10 else None

    # 计算RSI（简化版）
    def simple_rsi(prices, period=14):
        if len(prices) < period + 1:
            return None
        deltas = np.diff(prices[-period-1:])
        gains = deltas[deltas > 0].sum() / period
        losses = -deltas[deltas < 0].sum() / period
        if losses == 0:
            return 100
        rs = gains / losses
        return 100 - (100 / (1 + rs))

    return {
        "MA5": round(sma_5, 2) if sma_5 else None,
        "MA10": round(sma_10, 2) if sma_10 else None,
        "RSI14": round(simple_rsi(close_prices), 2) if simple_rsi(close_prices) is not None else None,
        "当前价格": close_prices[-1],
        "近期高点": max(close_prices[-10:]) if len(close_prices) >= 10 else None,
        "近期低点": min(close_prices[-10:]) if len(close_prices) >= 10 else None
    }

def create_visualization_in_memory(hist_df, stock_name, symbol):
    """创建数据可视化图表并返回内存中的图片对象"""
    try:
        # 设置中文字体 - 使用更通用的方法
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        
        # 创建子图
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(12, 10))
        
        # 转换日期格式以便绘图
        hist_df['日期'] = pd.to_datetime(hist_df['日期'])
        
        # 绘制价格走势
        ax1.plot(hist_df['日期'], hist_df['收盘'], label='收盘价', marker='o', linewidth=2)
        ax1.plot(hist_df['日期'], hist_df['开盘'], label='开盘价', marker='s', linewidth=1)
        ax1.set_title(f'{stock_name}({symbol}) - 价格走势', fontsize=14)
        ax1.set_ylabel('价格 (元)', fontsize=12)
        ax1.legend()
        ax1.grid(True, linestyle='--', alpha=0.6)
        
        # 绘制成交量
        ax2.bar(hist_df['日期'], hist_df['成交量'], alpha=0.6, color='orange', label='成交量')
        ax2.set_title('成交量', fontsize=14)
        ax2.set_ylabel('成交量', fontsize=12)
        ax2.set_xlabel('日期', fontsize=12)
        ax2.legend()
        ax2.grid(True, linestyle='--', alpha=0.6)
        
        # 旋转x轴标签以便更好地显示
        plt.xticks(rotation=45)
        plt.tight_layout()
        
        # 将图片保存到内存中
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()  # 关闭图形以释放内存
        
        return img_buffer
    except Exception as e:
        print(f"⚠️ 生成图表时出错: {str(e)}")
        return None

def call_llm_for_analysis(prompt):
    """调用大模型API生成分析报告"""
    import requests

    # 检查是否已配置API密钥
    if QWEN_API_KEY == "your_api_key_here":
        print("⚠️ 提示: 未配置API密钥，使用模拟AI分析。")
        return f"模拟AI分析: \n{prompt[:500]}...\n\n⚠️ 请在代码中配置真实的API密钥以获得完整AI分析报告。"
    
    headers = {
        "Authorization": f"Bearer {QWEN_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "qwen-turbo",  # 使用qwen-turbo模型，也可以使用qwen-plus或qwen-max
        "input": {
            "messages": [
                {"role": "system", "content": "你是一位资深股票分析师，擅长用简洁、专业的语言分析股票数据。"},
                {"role": "user", "content": prompt}
            ]
        },
        "parameters": {
            "temperature": 0.3,  # 降低随机性，使分析更稳定
        }
    }

    try:
        response = requests.post(QWEN_API_URL, json=payload, headers=headers, timeout=30)
        
        # 检查HTTP状态码
        if response.status_code != 200:
            print(f"⚠️ API请求失败，状态码: {response.status_code}")
            print(f"响应内容: {response.text[:200]}...")
            return f"API请求失败: HTTP {response.status_code}\n{response.text[:500]}"
        
        result = response.json()
        
        # 检查响应格式并提取内容
        if 'output' in result and 'text' in result['output']:
            return result['output']['text']
        elif 'error' in result:
            print(f"⚠️ API返回错误: {result['error']}")
            return f"API错误: {result['error']}"
        else:
            print(f"⚠️ API响应格式不正确: {result}")
            return f"API响应格式错误: {result}"
            
    except Exception as e:
        print(f"⚠️ AI模型调用失败: {str(e)}")
        return f"AI模型调用失败: {str(e)}\n以下是股票数据摘要:\n\n[实时行情]\n{prompt.split('【股票基本信息】')[1].split('【技术指标分析】')[0] if '【股票基本信息】' in prompt and '【技术指标分析】' in prompt else prompt[:1000]}"

def build_analysis_prompt(stock_data, tech_indicators, user_question):
    """构建给大模型的提示词"""

    prompt = f"""
请基于以下股票数据，生成一份简要分析报告，并直接回答用户问题。

【股票基本信息】
股票名称：{stock_data['实时行情']['名称']}
最新价格：{stock_data['实时行情']['最新价']}元
今日涨跌幅：{stock_data['实时行情']['涨跌幅']}%
成交量：{stock_data['实时行情']['成交量']:,}手

【技术指标分析】
{json.dumps(tech_indicators, indent=2, ensure_ascii=False)}

【近期股价历史（最近5天）】
{json.dumps(stock_data['历史数据'], indent=2, ensure_ascii=False)}

【用户具体问题】
"{user_question}"

请按以下结构组织回答：
1. 行情速览：总结当前股价位置和短期趋势
2. 技术面分析：基于指标分析多空力量
3. 关键提示：指出1-2个最重要的观察点
4. 风险提示：必含"投资有风险，决策需谨慎"
请使用专业但易懂的语言，避免过度复杂的术语。
"""
    return prompt

def create_word_report(stock_data, analysis_report, chart_buffer, user_question, symbol):
    """创建Word报告文档"""
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(f'股票分析报告 - {stock_data["实时行情"]["名称"]} ({symbol})', 0)
        
        # 添加分析时间
        doc.add_paragraph(f'分析时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'用户问题: {user_question}')
        
        # 添加股票基本信息
        doc.add_heading('股票基本信息', level=1)
        doc.add_paragraph(f'股票名称: {stock_data["实时行情"]["名称"]}')
        doc.add_paragraph(f'当前价格: {stock_data["实时行情"]["最新价"]}元')
        doc.add_paragraph(f'涨跌幅: {stock_data["实时行情"]["涨跌幅"]}%')
        doc.add_paragraph(f'成交量: {stock_data["实时行情"]["成交量"]:,} 手')
        
        # 添加技术指标
        doc.add_heading('技术指标分析', level=1)
        # 这里可以添加具体的技术指标分析
        
        # 添加AI分析报告
        doc.add_heading('AI分析报告', level=1)
        doc.add_paragraph(analysis_report)
        
        # 添加图表
        if chart_buffer:
            doc.add_heading('股价走势图', level=1)
            doc.add_paragraph('以下是该股票最近30天的价格走势和成交量图:')
            doc.add_picture(chart_buffer, width=Inches(6))  # 设置图片宽度为6英寸
        
        # 添加风险提示
        doc.add_heading('风险提示', level=1)
        doc.add_paragraph('投资有风险，决策需谨慎。以上分析仅供参考，不构成投资建议。')
        
        # 生成文件名，保存到脚本所在目录
        filename = f"{symbol}_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
        
        # 保存文档
        doc.save(filepath)
        
        return filename
    except Exception as e:
        print(f"⚠️ 生成Word报告时出错: {str(e)}")
        return None

def interactive_mode():
    """交互式模式：持续对话并生成Word报告"""
    print("="*60)
    print("🔍 欢迎使用股票分析AI智能体（交互式Word报告版）")
    print("💡 提示：输入 'quit' 或 'exit' 退出程序")
    print("💡 提示：输入 'change' + 股票代码 可切换分析的股票")
    print("💡 提示：每次提问后会自动生成Word报告")
    print("="*60)
    
    current_symbol = None
    current_stock_data = None
    current_hist_df = None
    
    while True:
        if current_symbol is None:
            # 首次运行，需要输入股票代码
            user_input = input("\n请输入股票代码（如 000001）: ").strip()
            if user_input.lower() in ['quit', 'exit', '退出']:
                print("👋 感谢使用，再见！")
                break
                
            current_symbol = user_input
            print(f"\n🔍 正在分析 [{current_symbol}]，请稍候...")
            print("-" * 50)

            # 1. 获取数据
            stock_data, msg = get_stock_data(current_symbol)
            if not stock_data:
                print(f"❌ 错误: {msg}")
                current_symbol = None  # 重置，重新输入
                continue

            # 获取完整的历史数据用于可视化
            end_date = datetime.now().strftime("%Y%m%d")
            start_date = (datetime.now() - timedelta(days=30)).strftime("%Y%m%d")
            current_hist_df = ak.stock_zh_a_hist(symbol=current_symbol, period="daily", start_date=start_date, end_date=end_date, adjust="qfq")

            current_stock_data = stock_data
            print(f"✅ 数据获取成功！当前分析股票: {stock_data['实时行情']['名称']}")
            
            # 询问用户问题
            question = input(f"\n请输入您想了解的问题（或输入 'help' 查看示例）: ").strip()
            if question.lower() == 'help':
                print("💡 示例问题：")
                print("  • 请分析这只股票的近期走势")
                print("  • 这只股票的技术指标如何？")
                print("  • 当前是否适合买入？")
                print("  • 请预测这只股票的短期趋势")
                question = input("\n请输入您的问题: ").strip()
        else:
            # 非首次运行，可以继续提问或切换股票
            user_input = input(f"\n[{current_symbol} {current_stock_data['实时行情']['名称']}] 请输入您的问题（'change'切换股票/'quit'退出）: ").strip()
            
            if user_input.lower() in ['quit', 'exit', '退出']:
                print("👋 感谢使用，再见！")
                break
            elif user_input.lower().startswith('change'):
                # 切换股票
                parts = user_input.split()
                if len(parts) >= 2:
                    new_symbol = parts[1]
                    print(f"\n🔍 正在切换到分析 [{new_symbol}]，请稍候...")
                    print("-" * 50)

                    stock_data, msg = get_stock_data(new_symbol)
                    if not stock_data:
                        print(f"❌ 错误: {msg}")
                        continue

                    # 获取新股票的历史数据
                    end_date = datetime.now().strftime("%Y%m%d")
                    start_date = (datetime.now() - timedelta(days=30)).strftime("%Y%m%d")
                    current_hist_df = ak.stock_zh_a_hist(symbol=new_symbol, period="daily", start_date=start_date, end_date=end_date, adjust="qfq")

                    current_symbol = new_symbol
                    current_stock_data = stock_data
                    print(f"✅ 股票切换成功！当前分析股票: {stock_data['实时行情']['名称']}")
                    
                    question = input(f"\n请输入您想了解的问题: ").strip()
                else:
                    print("❌ 请指定股票代码，例如: change 000001")
                    continue
            else:
                question = user_input
        
        if current_stock_data is None or current_hist_df is None:
            continue
            
        # 计算技术指标
        tech_indicators = calculate_technical_indicators(current_stock_data['历史数据'])

        # 生成可视化图表
        print("📊 正在生成可视化图表...")
        chart_buffer = create_visualization_in_memory(current_hist_df, current_stock_data['实时行情']['名称'], current_symbol)
        if chart_buffer:
            print("📈 图表生成成功")
        else:
            print("⚠️ 图表生成失败，继续执行分析...")

        # 构建提示词并调用AI
        print("🤖 正在生成AI分析报告...")
        prompt = build_analysis_prompt(current_stock_data, tech_indicators, question)
        analysis_report = call_llm_for_analysis(prompt)

        # 创建Word报告
        print("📝 正在生成Word报告...")
        word_filename = create_word_report(current_stock_data, analysis_report, chart_buffer, question, current_symbol)
        
        if word_filename:
            print(f"📄 Word报告已保存至: {word_filename}")
        else:
            print("⚠️ Word报告生成失败")

        # 输出结果到控制台
        print(f"\n📈 【{current_stock_data['实时行情']['名称']}】分析报告")
        print(f"📅 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"💰 当前价格: {current_stock_data['实时行情']['最新价']}元 ({current_stock_data['实时行情']['涨跌幅']}%)")
        if word_filename:
            print(f"📄 Word报告: {word_filename}")
        print("=" * 50)
        print(analysis_report)
        print("=" * 50)

def main():
    """主函数：处理命令行交互"""
    parser = argparse.ArgumentParser(description="股票分析AI智能体（交互式Word报告生成）")
    parser.add_argument("--interactive", "-i", action="store_true", help="交互式模式，可连续提问并生成Word报告")
    parser.add_argument("symbol", nargs='?', help="股票代码（如：000001、sz000001）")
    parser.add_argument("-q", "--question", default="请分析这只股票", help="你的具体问题")

    args = parser.parse_args()

    if args.interactive or args.symbol is None:
        # 进入交互模式
        interactive_mode()
    else:
        # 传统单次模式
        print(f"\n🔍 正在分析 [{args.symbol}]，请稍候...")
        print("-" * 50)

        # 1. 获取数据
        stock_data, msg = get_stock_data(args.symbol)
        if not stock_data:
            print(f"❌ 错误: {msg}")
            return

        # 2. 获取完整的历史数据用于可视化
        end_date = datetime.now().strftime("%Y%m%d")
        start_date = (datetime.now() - timedelta(days=30)).strftime("%Y%m%d")
        hist_df = ak.stock_zh_a_hist(symbol=args.symbol, period="daily", start_date=start_date, end_date=end_date, adjust="qfq")

        # 3. 计算技术指标
        tech_indicators = calculate_technical_indicators(stock_data['历史数据'])

        # 4. 生成可视化图表
        print("📊 正在生成可视化图表...")
        chart_buffer = create_visualization_in_memory(hist_df, stock_data['实时行情']['名称'], args.symbol)
        if chart_buffer:
            print("📈 图表生成成功")
        else:
            print("⚠️ 图表生成失败，继续执行分析...")

        # 5. 构建提示词并调用AI
        print("🤖 数据获取成功，正在生成AI分析报告...")
        prompt = build_analysis_prompt(stock_data, tech_indicators, args.question)
        analysis_report = call_llm_for_analysis(prompt)

        # 6. 创建Word报告
        print("📝 正在生成Word报告...")
        word_filename = create_word_report(stock_data, analysis_report, chart_buffer, args.question, args.symbol)
        
        if word_filename:
            print(f"📄 Word报告已保存至: {word_filename}")
        else:
            print("⚠️ Word报告生成失败")

        # 7. 输出结果到控制台
        print(f"\n📈 【{stock_data['实时行情']['名称']}】分析报告")
        print(f"📅 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"💰 当前价格: {stock_data['实时行情']['最新价']}元 ({stock_data['实时行情']['涨跌幅']}%)")
        if word_filename:
            print(f"📄 Word报告: {word_filename}")
        print("=" * 50)
        print(analysis_report)
        print("=" * 50)

if __name__ == "__main__":
    main()